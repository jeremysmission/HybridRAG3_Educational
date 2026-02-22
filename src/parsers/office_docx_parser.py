# ============================================================================
# HybridRAG -- DOCX Parser (src/parsers/office_docx_parser.py)
# ============================================================================
#
# WHAT THIS FILE DOES (plain English):
#   Reads Microsoft Word (.docx) files and extracts all the text from them.
#   A .docx file is actually a ZIP archive containing XML files. The
#   python-docx library knows how to unpack that ZIP and pull out
#   paragraphs of text, ignoring formatting (bold, font size, etc.)
#   that doesn't matter for search.
#
# WHY THIS MATTERS:
#   Many engineering reports and procedures are written in Word.
#   To search them, we need to extract the words into plain text
#   so they can be chunked, embedded, and indexed like any other file.
#
# HOW IT WORKS:
#   1. Open the .docx file using the python-docx library
#   2. Loop through every paragraph in the document
#   3. Skip empty paragraphs (blank lines, spacers)
#   4. Join all non-empty paragraphs with double newlines
#   5. Return the combined text + details (character count, paragraph count)
#
# ERROR HANDLING:
#   If the .docx file is corrupted, password-protected, or the python-docx
#   library is not installed, the parser returns empty text with an error
#   message in the details dict. It never crashes the indexing pipeline.
#
# INTERNET ACCESS: NONE
# ============================================================================

from __future__ import annotations

from pathlib import Path
from typing import Tuple, Dict, Any


class DocxParser:
    def parse(self, file_path: str) -> str:
        text, _ = self.parse_with_details(file_path)
        return text

    def parse_with_details(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract all paragraph text from a Word document.

        Returns (text, details_dict) where text is the full extracted
        content and details_dict has diagnostic info like character count.
        """
        path = Path(file_path)
        details: Dict[str, Any] = {"file": str(path), "parser": "DocxParser"}

        # python-docx is imported here (not at the top of the file) so that
        # the rest of HybridRAG still works even if python-docx is not installed.
        try:
            from docx import Document  # python-docx
        except Exception as e:
            details["error"] = f"IMPORT_ERROR: {type(e).__name__}: {e}"
            return "", details

        try:
            # Open the .docx file (internally unzips and parses the XML)
            doc = Document(str(path))

            # Walk through every paragraph and collect the non-empty ones.
            # A "paragraph" in Word includes headings, bullet points, and
            # regular body text -- we treat them all the same.
            parts = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)

            # Join paragraphs with blank lines between them (mimics the
            # visual spacing you see in the Word document).
            full = "\n\n".join(parts).strip()
            details["total_len"] = len(full)
            details["paragraphs"] = len(doc.paragraphs)

            return full, details
        except Exception as e:
            # Corrupted file, password-protected, or other issue --
            # return empty text so the indexer can skip this file gracefully.
            details["error"] = f"RUNTIME_ERROR: {type(e).__name__}: {e}"
            return "", details
